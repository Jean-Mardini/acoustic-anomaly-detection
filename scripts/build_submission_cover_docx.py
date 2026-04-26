"""One-off: generate Word cover sheet for course submission. Requires: pip install python-docx"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parents[1] / "deliverables" / "Course_Submission_Cover_Sheet.docx"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    d = Document()
    t = d.add_paragraph("USJ — Machine Learning (final project)")
    t.runs[0].font.size = Pt(12)
    d.add_heading("Acoustic Anomaly Detection — submission links", level=0)

    d.add_heading("Group", level=1)
    for name in (
        "Jean Mardini",
        "Marc Khattar",
        "Christy Tannoury",
        "Angela Nabhan",
    ):
        d.add_paragraph(name, style="List Bullet")

    d.add_heading("1. GitHub repository (code, README, optional files)", level=1)
    p = d.add_paragraph()
    p.add_run("Repository: ").bold = True
    p.add_run("https://github.com/Jean-Mardini/acoustic-anomaly-detection")
    d.add_paragraph("Clone: git clone https://github.com/Jean-Mardini/acoustic-anomaly-detection.git")

    d.add_heading("2. Public Docker image (required)", level=1)
    p = d.add_paragraph()
    p.add_run("Image page (use this in reports): ").bold = True
    p.add_run("https://hub.docker.com/r/jeanmardini/acoustic-anomaly-detection")
    d.add_paragraph(
        "Alternative (repository settings / same image): "
        "https://hub.docker.com/repository/docker/jeanmardini/acoustic-anomaly-detection/general"
    )
    p = d.add_paragraph()
    p.add_run("Pull: ").bold = True
    p.add_run("docker pull jeanmardini/acoustic-anomaly-detection:latest")
    p = d.add_paragraph()
    p.add_run("Run: ").bold = True
    p.add_run("docker run --rm -p 8000:8000 jeanmardini/acoustic-anomaly-detection:latest")
    d.add_paragraph("UI: http://localhost:8000  ·  API: http://localhost:8000/docs")

    d.add_heading("3. Files submitted with this package", level=1)
    d.add_paragraph("Adjust names if you zip different filenames.", style="Intense Quote")
    items = [
        "Video demo (≈3 min): in the attached .zip (or a single .mp4 in the zip).",
        "Presentation: acoustic_anomaly_detection.pptx (also on GitHub in repo root).",
        "One-page / project report: project_report.docx (also on GitHub in repo root).",
    ]
    for x in items:
        d.add_paragraph(x, style="List Number")

    d.add_heading("4. What the grader can do", level=1)
    d.add_paragraph(
        "Open the GitHub link for the full project; use the Docker Hub link to verify the "
        "public image; watch the video from the zip; open the PowerPoint and report from this "
        "folder or from the repository."
    )

    d.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
